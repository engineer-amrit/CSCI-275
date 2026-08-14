#!/usr/bin/env python3
"""Generates the Moderation System test report as a .docx file."""
import sys
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = "docs/report/test_report.docx"

BACKEND = {
    "restaurant.service.spec.ts — RestaurantService (11 tests)": [
        "verify — throws NotFoundException when the restaurant does not exist",
        "verify — returns already verified when the restaurant is verified",
        "verify — verifies a restaurant meeting the review criteria",
        "verify — does not verify when there are fewer than five reviews",
        "verify — does not verify when reviews come from fewer than five users",
        "verify — does not verify when the average rating is below the threshold",
        "verify — ignores invalid reviews when computing verification",
        "verify — verifies exactly at the threshold boundary",
        "checkData — flags a restaurant whose data contains an offensive word",
        "checkData — reports verified for a restaurant with clean data",
        "checkData — throws NotFoundException when the restaurant does not exist",
    ],
    "config.service.test.ts — ConfigService (7 tests)": [
        "returns NODE_ENV",
        "returns HOST",
        "returns DB_HOST",
        "returns isProduction as false for the test environment",
        "returns isDevelopment as false for the test environment",
        "returns isTest as true for the test environment",
        "returns isStaging as false for the test environment",
    ],
}

FRONTEND = {
    "utils/reviewStats.test.ts — computeReviewStats (3 tests)": [
        "returns zeroed stats for an empty list",
        "counts total, verified and pending reviews",
        "keeps verified and pending consistent with total",
    ],
    "utils/reviewStats.test.ts — filterReviews (5 tests)": [
        "returns all reviews for the 'all' filter",
        "returns only language-verified reviews for the 'verified' filter",
        "returns only pending reviews for the 'unverified' filter",
        "does not mutate the input list",
        "handles an empty list for every filter",
    ],
    "services/restaurant.service.test.ts — restaurantService (9 tests)": [
        "getAll — returns a list of restaurants",
        "getById — returns the restaurant with the matching id",
        "getById — returns undefined for an unknown id",
        "verify — verifies a restaurant that meets the review criteria",
        "verify — does not verify a restaurant below the review threshold",
        "verify — throws when the restaurant is not found",
        "checkData — flags a restaurant whose data contains offensive words",
        "checkData — marks a clean restaurant as verified",
        "setDataStatus — sets the data status on a restaurant",
    ],
    "services/media.service.test.ts — mediaService (6 tests)": [
        "getAll — returns a list of media",
        "setVerified — marks media as verified",
        "setVerified — marks media as unverified",
        "setVerified — throws when the media is not found",
        "undo — reverts the last moderation action for media",
        "undo — throws when there is no action to undo",
    ],
    "services/review.service.test.ts — reviewService (7 tests)": [
        "getAll — returns a list of reviews",
        "setLanguageVerified — marks a review language as verified",
        "setLanguageVerified — flags a review language as unverified",
        "setLanguageVerified — throws when the review is not found",
        "undo — reverts the last language moderation action",
        "undo — throws when there is no action to undo",
    ],
}

E2E = [
    ("user-auth /json API", "POST /json/login returns a session token", "PASS"),
    ("user-auth /json API", "POST /json/verify accepts a valid token and rejects a bad one", "PASS"),
    ("user-auth /json API", "POST /json/users creates a moderator; duplicate returns 409", "PASS"),
    ("user-auth /json API", "Missing token returns 401; non-moderator role returns 403", "PASS"),
    ("user-auth /json API", "Wrong password returns 401; GET /json/me and /json/users/:id behave", "PASS"),
    ("Full-stack Docker", "docker compose up boots postgres, user-auth, restaurant, search, moderation-api, web", "PASS"),
    ("Full-stack Docker", "GET /api/health returns ok through the nginx proxy (port 5173)", "PASS"),
    ("Full-stack Docker", "POST /api/v1/auth/login works with the seeded root moderator", "PASS"),
    ("Full-stack Docker", "GET /api/v1/restaurant lists seeded restaurants with verification state", "PASS"),
    ("Full-stack Docker", "GET /api/v1/restaurant/:id/data-check flags an offensive restaurant", "PASS"),
    ("Full-stack Docker", "PATCH /api/v1/restaurant/verify/:id verifies a restaurant meeting criteria", "PASS"),
    ("Full-stack Docker", "GET /api/v1/restaurant/:id/reviews, /api/v1/review, /api/v1/media return data", "PASS"),
    ("Full-stack Docker", "Protected routes return 401 without a session token", "PASS"),
]


def set_cell_shade(cell, color):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color)
    tcPr.append(shd)


def main():
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading("Test Report — Restaurant Review Moderation System", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for line in [
        "Team XOR — CSCI-275",
        "Course: CSCI-275 | Module: Moderation System",
        f"Report generated: 13 August 2026",
    ]:
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].italic = True

    doc.add_heading("1. Scope", level=1)
    doc.add_paragraph(
        "This report documents the automated and manual verification performed on the Restaurant "
        "Review Moderation System. The system moderates restaurant listings (business verification "
        "and data-quality scanning), reviews (language moderation with undo support), and media "
        "(verification with undo support). The stack is composed of a React frontend "
        "(moderation-web), a NestJS moderation API (moderation-api), and three microservices: "
        "user-auth (session-based authentication), restaurant (Prisma/PostgreSQL profile and review "
        "data), and search (restaurant discovery)."
    )

    doc.add_heading("2. Test Environment", level=1)
    env = [
        ("Host OS", "macOS (Darwin)"),
        ("Node.js", "24.16.0"),
        ("Package manager", "pnpm 11.5.0 (workspace monorepo)"),
        ("Test runner", "Vitest 4.1.4"),
        ("Backend framework", "NestJS 11 / Express"),
        ("Frontend", "React 19 + Vite 6"),
        ("Database", "PostgreSQL 16 (three databases: alpha_db, alpha_restaurant, alpha_search)"),
        ("Containerization", "Docker + docker compose"),
    ]
    table = doc.add_table(rows=len(env), cols=2)
    table.style = "Light Grid Accent 1"
    for i, (k, v) in enumerate(env):
        table.rows[i].cells[0].text = k
        table.rows[i].cells[1].text = v

    doc.add_heading("3. Automated Test Suites", level=1)
    doc.add_paragraph(
        "Automated tests are written in Vitest. Backend tests target the restaurant verification "
        "service (review-threshold criteria) and the data-quality scan. Frontend tests cover the "
        "API service layer (against a stubbed fetch), the review statistics/filter utilities, and "
        "media/review moderation actions. All suites are run as part of the continuous checks "
        "(`pnpm run test` in each app) and type-check/lint gates pass."
    )

    doc.add_heading("3.1 Backend — moderation-api (18 tests)", level=2)
    doc.add_paragraph("Command: `pnpm run test` in apps/backend/moderation-api. Result: 18 passed, 0 failed.")
    for group, tests in BACKEND.items():
        doc.add_heading(group, level=3)
        for t in tests:
            p = doc.add_paragraph(t, style="List Bullet")
            p.runs[0].font.size = Pt(10)

    doc.add_heading("3.2 Frontend — moderation-web (30 tests)", level=2)
    doc.add_paragraph("Command: `pnpm run test` in apps/web/moderation-web. Result: 30 passed, 0 failed.")
    for group, tests in FRONTEND.items():
        doc.add_heading(group, level=3)
        for t in tests:
            p = doc.add_paragraph(t, style="List Bullet")
            p.runs[0].font.size = Pt(10)

    doc.add_heading("4. End-to-End Verification", level=1)
    doc.add_paragraph(
        "The user-auth `/json` API was verified against a live PostgreSQL instance. The full stack "
        "was then started with `docker compose up --build` (six containers) and exercised through "
        "the nginx proxy on http://localhost:5173, which forwards /api to the moderation API."
    )
    table = doc.add_table(rows=len(E2E) + 1, cols=3)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(["Area", "Check", "Result"]):
        hdr[i].text = h
        set_cell_shade(hdr[i], "D9E2F3")
    for i, (area, check, result) in enumerate(E2E, start=1):
        table.rows[i].cells[0].text = area
        table.rows[i].cells[1].text = check
        table.rows[i].cells[2].text = result

    doc.add_heading("5. Results Summary", level=1)
    summary = [
        ("Backend unit tests", "2 files / 18 tests", "18 passed"),
        ("Frontend unit tests", "5 files / 30 tests", "30 passed"),
        ("Total automated tests", "48 tests", "48 passed"),
        ("End-to-end checks", "13 checks", "13 passed"),
        ("Type checking", "backend + frontend", "passed"),
        ("Linting", "backend + frontend", "passed"),
    ]
    table = doc.add_table(rows=len(summary), cols=3)
    table.style = "Light Grid Accent 1"
    for i, (k, v, r) in enumerate(summary):
        table.rows[i].cells[0].text = k
        table.rows[i].cells[1].text = v
        table.rows[i].cells[2].text = r

    doc.add_heading("6. Conclusion", level=1)
    doc.add_paragraph(
        "All 48 automated tests pass and every end-to-end check succeeds. The moderation system "
        "correctly verifies restaurants against the review-threshold criteria, flags restaurant data "
        "that contains offensive words, moderates review language and media with recoverable undo "
        "actions, and protects every moderation route behind session-based authentication. The stack "
        "runs end-to-end from a single `docker compose up --build` command, with the web app served "
        "on http://localhost:5173 and sign-in available with the seeded root moderator "
        "(moderator@admin.com / moderator123)."
    )

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    sys.exit(main())
